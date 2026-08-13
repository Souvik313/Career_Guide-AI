from pathlib import Path
import re
import fitz
from docx import Document
from io import BytesIO


class ResumeParser:

    def __init__(self):
        pass

    def extract_text(self, file_path):
        file_path = Path(file_path)

        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._extract_pdf(file_path)

        elif suffix == ".docx":
            return self._extract_docx(file_path)

        else:
            raise ValueError(
                "Unsupported file format. Only PDF and DOCX are supported."
            )

    def extract_text_from_bytes(
        self,
        file_bytes: bytes,
        filename: str,
    ):
        """
        Extract text directly from uploaded file bytes.

        This is used when files are stored in Cloudinary
        rather than on the local filesystem.
        """

        suffix = Path(filename).suffix.lower()

        if suffix == ".pdf":
            return self._extract_pdf_bytes(file_bytes)

        elif suffix == ".docx":
            return self._extract_docx_bytes(file_bytes)

        else:
            raise ValueError(
                "Unsupported file format. Only PDF and DOCX are supported."
            )

    def _extract_pdf_bytes(self, file_bytes: bytes):
        """
        Extract text from PDF bytes using PyMuPDF.
        """

        pages = []

        with fitz.open(
            stream=file_bytes,
            filetype="pdf",
        ) as document:

            for page in document:
                pages.append(page.get_text())

        return "\n".join(pages)

    def _extract_docx_bytes(self, file_bytes: bytes):
        """
        Extract text from DOCX bytes.
        """

        document = Document(
            BytesIO(file_bytes)
        )

        paragraphs = []

        for paragraph in document.paragraphs:
            paragraphs.append(paragraph.text)

        return "\n".join(paragraphs)

    def extract_candidate_name(self, resume_text):
        """
        Attempts to extract the candidate's full name
        from the first few lines of the resume.
        """

        lines = [
            line.strip()
            for line in resume_text.split("\n")
            if line.strip()
        ]

        for line in lines[:5]:

            if any(
                keyword in line.lower()
                for keyword in [
                    "@",
                    "linkedin",
                    "github",
                    "phone",
                    "email",
                    "resume",
                    "curriculum vitae",
                    "address",
                ]
            ):
                continue

            if re.fullmatch(
                r"[A-Za-z.\- ]{3,50}",
                line,
            ):

                words = line.split()

                if 2 <= len(words) <= 4:
                    return line.title()

        return "Unknown Candidate"

    def _extract_pdf(self, file_path):
        """
        Extract text from a PDF resume using PyMuPDF.
        """

        pages = []

        with fitz.open(file_path) as document:

            for page in document:
                pages.append(page.get_text())

        return "\n".join(pages)

    def _extract_docx(self, file_path):
        document = Document(file_path)

        paragraphs = []

        for paragraph in document.paragraphs:
            paragraphs.append(paragraph.text)

        return "\n".join(paragraphs)